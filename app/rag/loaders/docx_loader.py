from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from rag.loaders.base import RawDocument, DocumentLoadError, BaseLoader


class DocxLoader(BaseLoader):
    """Loads .docx files into a RawDocument, preserving paragraph and table content."""

    SUPPORTED_EXTENSIONS = (".docx",)

    def load(self, file_path: str | Path) -> RawDocument:
        path = Path(file_path)

        if not path.exists():
            raise DocumentLoadError(f"File not found: {path}")
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise DocumentLoadError(f"Unsupported file type for DocxLoader: {path.suffix}")

        try:
            doc = Document(str(path))
        except Exception as e:
            raise DocumentLoadError(f"Failed to open docx file {path}: {e}") from e

        try:
            full_text = self._extract_ordered_text(doc)
        except Exception as e:
            raise DocumentLoadError(f"Failed to parse docx content in {path}: {e}") from e

        if not full_text.strip():
            raise DocumentLoadError(f"No extractable text found in {path}")

        metadata = self._build_metadata(path, doc)

        return RawDocument(text=full_text, metadata=metadata)

    def _extract_ordered_text(self, doc: Document) -> str:
        """
        Walks the document body in order so paragraphs and tables appear
        interleaved as they do in the original file, rather than all
        paragraphs first and all tables after.
        """
        blocks: list[str] = []
        body = doc.element.body

        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]  # strip namespace

            if tag == "p":
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    blocks.append(text)

            elif tag == "tbl":
                table = Table(child, doc)
                table_text = self._extract_table_text(table)
                if table_text:
                    blocks.append(table_text)

        return "\n\n".join(blocks)

    def _extract_table_text(self, table: Table) -> str:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)

    def _build_metadata(self, path: Path, doc: Document) -> dict[str, Any]:
        core_props = doc.core_properties
        return {
            "source": str(path),
            "file_type": "docx",
            "file_name": path.name,
            "title": core_props.title or None,
            "author": core_props.author or None,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
        }